import os
import pytesseract
import pyttsx3
from PIL import Image
import pdfplumber
from docx import Document
from transformers import pipeline
import customtkinter
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
from tkinter import *
from tkinter import ttk
import tkinter as tk
from PIL import ImageTk, Image
from deep_translator import GoogleTranslator
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
#pip install python-docx
engine = pyttsx3.init()
import pytesseract

# Set the tesseract executable path (update it according to your system)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
def extract_text_from_image(file_path):
    try:
        image = Image.open(file_path)
        return pytesseract.image_to_string(image)
    except Exception as e:
        return f"Error reading image: {e}"

def extract_text_from_pdf(file_path):
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF: {e}"

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading DOCX: {e}"

# --- Summarization Function ---
def summarize_text(text, max_chunk=1000):
    chunks = [text[i:i + max_chunk] for i in range(0, len(text), max_chunk)]
    summary = ""
    for chunk in chunks:
        if len(chunk.strip()) > 50:
            summarized = summarizer(chunk, max_length=130, min_length=30, do_sample=False)[0]['summary_text']
            summary += summarized + "\n"
    return summary.strip()

# --- GUI File Handler ---
def generate_report():
    file_path = filedialog.askopenfilename(title="Select a file",
                                           filetypes=[("Documents", "*.pdf *.docx *.jpg *.jpeg *.png")])
    if not file_path:
        return

    ext = os.path.splitext(file_path)[1].lower()

    if ext in ['.png', '.jpg', '.jpeg']:
        text = extract_text_from_image(file_path)
    elif ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        messagebox.showerror("Unsupported Format", "Please select a PDF, DOCX, or image file.")
        return

    if not text.strip():
        messagebox.showinfo("No Text Found", "No readable text was extracted from the file.")
        return

    output_text.delete("1.0", tk.END)
    output_text.insert(tk.END, "Generating summary...\n\n")
    window.update()

    try:
        summary = summarize_text(text)
        output_text.delete("1.0", tk.END)
        output_text.insert(tk.END, summary)
    except Exception as e:
        output_text.delete("1.0", tk.END)
        output_text.insert(tk.END, f"Error generating summary: {e}")
def translate_summary():

    cont1 = output_text.get("1.0", tk.END)

    ll1=L.get()
    if ll1=='Kannada':
        ll2='kn'
    if ll1=='Hindi':
        ll2='hi'


    # Kannada text
    kannada_text = cont1

    # Translate to English
    translated_text = GoogleTranslator(source='en', target=ll2).translate(kannada_text)

    print("Original (English):", kannada_text)
    print("Translated (Kannad):", translated_text)

    Tran_text.insert(tk.END, translated_text)



    # Text-to-Speech
def speak_summary():
    content = output_text.get("1.0", tk.END)
    print(content)
    if content.strip():

        engine.say(content)
        engine.runAndWait()


def save():
    text = Tran_text.get("1.0", tk.END)
    with open("output.txt", "w", encoding="utf-8") as file:
        file.write(text.strip())
    text = output_text.get("1.0", tk.END)
    with open("output1.txt", "w", encoding="utf-8") as file:
        file.write(text.strip())
    from docx import Document

    text = Tran_text.get("1.0", tk.END)
    doc = Document()
    doc.add_paragraph(text.strip())
    doc.save("output.docx")
    text = output_text.get("1.0", tk.END)
    doc = Document()
    doc.add_paragraph(text.strip())
    doc.save("output1.docx")



def save1():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()

    # Add Unicode font
    pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
    pdf.set_font("DejaVu", size=12)

    text = Tran_text.get("1.0", tk.END).strip()

    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)

    pdf.output("output.pdf")
    pdf = FPDF()
    pdf.add_page()

    # Add Unicode font
    pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
    pdf.set_font("DejaVu", size=12)

    text = output_text.get("1.0", tk.END).strip()

    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)

    pdf.output("output1.pdf")




def clear():
    Tran_text.delete("1.0", tk.END)
    output_text.delete("1.0", tk.END)

window = tk.Tk()
window.title("NLP Report Generator")
window.geometry("1366x768")
window.configure(bg="#0abbf0")
canv = Canvas(window, width=1366, height=768, bg='white')
canv.grid(row=2, column=3)
img = Image.open('back3.png')
photo = ImageTk.PhotoImage(img)
canv.create_image(1,1, anchor=NW, image=photo)
L=StringVar()
l1 = []
def exit():
   window.destroy()
   os.system("python main.py")
def keyin():

    os.system("python QA.py")
L.set('Select Language')
l1.append("Kannada")
l1.append("Hindi")
t1=OptionMenu(window,L,*l1)
t1.config(bg="white")
t1["menu"].config(bg="white")
t1.place(x=690,y=150)
tk.Button(window, text="Key-Insights",width=18,height=1, bg="white",border=0,cursor="hand1",font=("Georgia", 12),command=keyin).place(x=100,y=150)

tk.Button(window, text="Select File & Generate Report", width=25,height=1,bg="white",border=0,cursor="hand1",font=("Georgia", 12),command=generate_report).place(x=370,y=150)

tk.Button(window, text="Translate",bg="white",width=13,height=1,border=0,font=("Georgia", 12),cursor="hand1", command=translate_summary).place(x=880,y=150)
tk.Button(window, text="Read Aloud",bg="white",width=8,height=1,border=0,font=("Georgia", 12),cursor="hand1", command=speak_summary).place(x=1060,y=150)
tk.Button(window, text=".docx",bg="white",height=1,border=0,font=("Georgia", 12),cursor="hand1", command=save).place(x=1070,y=230)
tk.Button(window, text=".pdf",bg="white",height=1,border=0,font=("Georgia", 12),cursor="hand1", command=save1).place(x=1150,y=230)
tk.Button(window, text="Clear",bg="white",width=8,height=1,border=0,cursor="hand1",font=("Georgia", 12), command=clear).place(x=1180,y=150)
tk.Button(window, text="Exit",bg="white",width=13,height=1,cursor="hand1",border=0,font=("Georgia", 12), command=exit).place(x=1200,y=650)
tk.Label(window, text="File Contents",bg="#03396f",fg="white",cursor="hand1",font=("Georgia", 12)).place(x=100,y=230)
output_text = scrolledtext.ScrolledText(window, wrap=tk.WORD, width=55, height=20, font=("Georgia", 10))
output_text.place(x=100,y=300)
tk.Label(window, text="Translated Contents",bg="#03396f",fg="white",font=("Georgia", 12)).place(x=700,y=230)
Tran_text = scrolledtext.ScrolledText(window, wrap=tk.WORD, width=60, height=20, font=("Georgia", 10))
Tran_text.place(x=700,y=300)
window.mainloop()
